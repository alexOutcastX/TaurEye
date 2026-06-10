#!/usr/bin/env python3
import pathlib
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = pathlib.Path("/home/user/TaurEye")
OUT = ROOT / "TaurEye-Where-We-Are.docx"

GREEN = RGBColor(0x12, 0x8F, 0x63)

doc = Document()

# ---- Title ----
t = doc.add_heading("TaurEye — Where We Are", level=0)
sub = doc.add_paragraph("Product status & competitive comparison")
sub.runs[0].italic = True
doc.add_paragraph("NSE/BSE end-of-day stock screener · web + Android (Capacitor) · static JSON data bundle + Supabase cloud layer.")

# ---- Section 1: status ----
doc.add_heading("1. Where We Are — detailed status", level=1)
status_rows = [
    ("Screening", "Filter-based screener", "Live", "Over the metrics bundle (~5,800 stocks)"),
    ("Screening", "Natural-language screener", "Live", "Plain-English → filters; signature feature"),
    ("Screening", "Saved screens", "Live", "Local; cloud table exists in schema"),
    ("Charts", "Native charts (lightweight-charts)", "Live", "Timeframes, moving averages"),
    ("Charts", "TradingView widget toggle", "Live", "TV's own data"),
    ("Charts", "Auto pattern detection", "Live", "Drawn as lines (patterns.ts)"),
    ("Charts", "Per-stock Key Data", "Live", "RSI, SMA50/200, 52w, ATR + SEBI disclaimer"),
    ("Data", "Global indices dashboard", "Live", "+ ticker"),
    ("Data", "Scrip search", "Live", ""),
    ("Engagement", "Watchlists", "Live", ""),
    ("Engagement", "EOD price alerts", "Live", "Above/below"),
    ("Engagement", "Push notifications (FCM)", "Live", "Native, via Firebase"),
    ("Engagement", "Light/dark theme", "Live", ""),
    ("Accounts", "Supabase auth (email + guest)", "Live", ""),
    ("Accounts", "Cloud credit wallet", "Live", "Server balance, unforgeable (RLS)"),
    ("Accounts", "Signup bonus (50) + daily claim (5)", "Live", "Secure RPCs"),
    ("Accounts", "Referral program", "Partial", "Tables + page exist; reward grant not wired"),
    ("AI / Reports", "AI stock analysis", "Built", "Edge Function (Haiku); needs deploy + ANTHROPIC_API_KEY"),
    ("AI / Reports", "Advanced report", "Built", "Credit sink"),
    ("Monetization", "Broker affiliate (Zerodha)", "LIVE & earning", "Angel/Upstox/Dhan = paste links"),
    ("Monetization", "Ads — AdSense (web)", "Built, gated", "Needs ca-pub-… + unit IDs"),
    ("Monetization", "Ads — AdMob (app)", "Built, gated", "Needs IDs + APK rebuild + privacy policy"),
    ("Monetization", "Razorpay credit purchases", "Built", "Needs Razorpay account + webhook"),
    ("Monetization", "Credit charging", "Off (by design)", "Flip after AI + purchases live"),
    ("Monetization", "Pro subscription", "Planned", "Phase D (Rs 999/yr)"),
    ("Platforms", "Web (nginx, static)", "Live", "Deploys from main in ~1 min"),
    ("Platforms", "Android APK", "Builds", "GitHub Actions"),
    ("Platforms", "OTA updates (Capgo)", "Wired", "Needs CAPGO_TOKEN + CAPGO_ENABLED"),
    ("Platforms", "iOS", "Scaffolded", "Not a current build target"),
    ("Infra", "Nightly data pipeline", "Live", "EOD → CA-adjust → JSON export"),
    ("Infra", "Compliance (SEBI/RLS/secrets)", "In place", "Factual-only, keys server-side"),
]
tbl = doc.add_table(rows=1, cols=4)
tbl.style = "Light Grid Accent 1"
for i, h in enumerate(["Area", "Feature", "Status", "Notes"]):
    c = tbl.rows[0].cells[i]
    c.text = h
    c.paragraphs[0].runs[0].bold = True
for area, feat, st, note in status_rows:
    cells = tbl.add_row().cells
    cells[0].text, cells[1].text, cells[2].text, cells[3].text = area, feat, st, note
    r = cells[2].paragraphs[0].runs
    if r:
        r[0].bold = True
        if "LIVE" in st or st == "Live":
            r[0].font.color.rgb = GREEN

# ---- Section 2: comparison ----
doc.add_heading("2. Comparison with peers", level=1)
comp_header = ["Capability", "TaurEye", "screener.in", "Chartink", "Trendlyne", "TradingView"]
comp_rows = [
    ("Primary focus", "Unified EOD screener (fund + tech)", "Fundamentals", "Technical scans", "Everything (fund+tech+analyst)", "Charting platform"),
    ("Fundamentals depth", "Key ratios (growing)", "Deep (10-yr) ★", "Minimal", "Broad + DVM scores", "Weak (India)"),
    ("Technical screening", "Yes", "Weak", "Strong ★", "Yes", "Generic"),
    ("Intraday / real-time", "No (EOD only)", "No", "Yes (near real-time)", "Partial", "Yes ★"),
    ("NL / plain-English screen", "Yes ★", "No (query lang)", "No (own syntax)", "No", "No (Pine Script)"),
    ("Charting", "Good", "Basic", "Basic", "Decent", "World-class ★"),
    ("Auto pattern detection", "Yes (drawn for you)", "No", "Via scans", "Some", "Via Pine/indicators"),
    ("Alerts", "Yes (EOD + push)", "Limited", "Yes", "Yes (paywalled)", "Yes"),
    ("AI stock analysis", "Yes (built-in, Haiku)", "No", "No", "Some", "Limited"),
    ("Mobile app quality", "Mobile-first native ★", "Mediocre", "Mediocre", "OK", "Good"),
    ("India focus", "Yes", "Yes", "Yes", "Yes", "No"),
    ("Price / year", "Free / Rs 999 Pro", "~Rs 5,000", "~Rs 2,400", "Rs 999–6,000", "~Rs 12,000+ (INR)"),
    ("Free tier", "Generous ★", "Limited", "Yes (ads)", "Heavily paywalled", "Limited"),
]
tbl2 = doc.add_table(rows=1, cols=len(comp_header))
tbl2.style = "Light Grid Accent 1"
for i, h in enumerate(comp_header):
    c = tbl2.rows[0].cells[i]
    c.text = h
    run = c.paragraphs[0].runs[0]
    run.bold = True
    if h == "TaurEye":
        run.font.color.rgb = GREEN
for row in comp_rows:
    cells = tbl2.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val
        if i == 0:
            cells[i].paragraphs[0].runs[0].bold = True
        if i == 1 and cells[i].paragraphs[0].runs:  # TaurEye column
            cells[i].paragraphs[0].runs[0].font.color.rgb = GREEN

p = doc.add_paragraph()
p.add_run("Our wedge: ").bold = True
p.add_run("we don't beat any one of them at their core strength — instead we bundle fundamentals + "
          "technicals + plain-English search + AI + patterns + alerts into one fast, mobile-first app, "
          "free to start, undercutting a field where everyone else charges Rs 2.4k–12k and silos their features.")

doc.save(str(OUT))
print("WROTE", OUT)
